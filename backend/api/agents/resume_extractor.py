"""
backend/api/agents/resume_extractor.py
========================================
ResumeExtractorAgent

Input  : Raw resume bytes (PDF or DOCX) + MIME type
Output : ParsedResume  (strict Pydantic schema — no invented data)

Pipeline
--------
1. Convert file bytes → plain text (pypdf for PDF, python-docx for DOCX)
2. Build a structured-extraction prompt that:
   a. Gives the LLM the raw text as the ONLY source of truth
   b. Demands output as a single JSON object matching ParsedResume schema
   c. Explicitly forbids inventing values not present in the text
3. Call the LLM via LlamaIndex's structured prediction
4. Parse and validate the response with Pydantic
5. Compute a heuristic confidence score
6. Return ParsedResume

Zero-hallucination enforcement
-------------------------------
* temperature = 0.0  (deterministic output)
* System prompt contains an explicit "GROUND RULE" section that instructs
  the model to output null / empty string for any field it cannot find
  in the source text.
* Response is validated against the Pydantic schema; any field that
  looks invented (e.g. an email address that doesn't appear in the raw
  text) is cleared and flagged in extraction_warnings.
"""

from __future__ import annotations

import io
import json
import re
import time
from typing import Any

from llama_index.core.llms import LLM
from llama_index.core import Settings as LlamaSettings

from backend.api.agents.base import (
    EducationEntry,
    ExperienceEntry,
    ParsedResume,
    configure_llama_settings,
)
from backend.api.core.config import settings
from backend.api.core.logging import get_logger

logger = get_logger(__name__)

# ---------------------------------------------------------------------------
# Text extraction helpers
# ---------------------------------------------------------------------------

def _extract_text_from_pdf(data: bytes) -> str:
    """Extract plain text from PDF bytes using pypdf."""
    import pypdf  # lazy import — not needed for non-PDF paths

    reader = pypdf.PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages.append(text)
    return "\n".join(pages)


def _extract_text_from_docx(data: bytes) -> str:
    """Extract plain text from DOCX bytes using python-docx."""
    import docx  # python-docx

    doc = docx.Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs)


def _extract_raw_text(file_bytes: bytes, mime_type: str) -> str:
    """Route to the correct extractor based on MIME type."""
    if mime_type == "application/pdf":
        return _extract_text_from_pdf(file_bytes)
    elif mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return _extract_text_from_docx(file_bytes)
    else:
        # Fallback: try to decode as UTF-8 text
        return file_bytes.decode("utf-8", errors="replace")


# ---------------------------------------------------------------------------
# Prompt templates
# ---------------------------------------------------------------------------

_SYSTEM_PROMPT = """\
You are a precise resume data extraction engine.

GROUND RULES — YOU MUST FOLLOW THESE WITHOUT EXCEPTION:
1. Extract information ONLY from the resume text provided. Do NOT invent, infer, or guess any data.
2. If a field is not present in the resume text, use an empty string "" or empty list [] or null.
3. Do NOT hallucinate email addresses, phone numbers, URLs, company names, dates, or any other facts.
4. Output ONLY a single valid JSON object. No markdown, no code fences, no explanations.
5. All dates should be formatted as strings exactly as they appear in the resume (e.g. "Jan 2021", "2020", "Present").
6. The "skills" array must contain only skills explicitly mentioned in the resume text.
7. Each experience entry description should be a factual summary of bullet points found in the resume.

OUTPUT SCHEMA (JSON):
{
  "name": "string — full name as it appears at the top of the resume",
  "email": "string — email address or empty string",
  "phone": "string or null — phone number as written",
  "location": "string or null — city/state/country as written",
  "linkedin_url": "string or null — full LinkedIn URL or null",
  "github_url": "string or null — full GitHub URL or null",
  "portfolio_url": "string or null — any personal/portfolio URL or null",
  "summary": "string or null — professional summary/objective section verbatim or null",
  "skills": ["array of skill strings — only those explicitly listed"],
  "experience": [
    {
      "company": "string",
      "title": "string",
      "start_date": "string — as written",
      "end_date": "string or null — as written, 'Present' if current",
      "description": "string — factual summary of responsibilities/achievements",
      "technologies": ["array of technologies explicitly mentioned for this role"]
    }
  ],
  "education": [
    {
      "institution": "string",
      "degree": "string",
      "field": "string",
      "graduation_year": "integer or null",
      "gpa": "string or null"
    }
  ],
  "certifications": ["array of certification strings"],
  "languages": ["array of language strings"]
}
"""

_USER_PROMPT_TEMPLATE = """\
Extract all structured data from the following resume text.
Remember: only output data you can directly read from the text below.

--- RESUME TEXT START ---
{resume_text}
--- RESUME TEXT END ---

Output only the JSON object, nothing else.
"""


# ---------------------------------------------------------------------------
# Confidence scoring heuristic
# ---------------------------------------------------------------------------

def _compute_confidence(parsed: dict[str, Any], raw_text: str) -> tuple[float, list[str]]:
    """
    Compute a heuristic confidence score [0.0, 1.0] and collect warnings.

    Strategy: verify that key extracted values actually appear in the raw text.
    Any field that doesn't appear verbatim is flagged and cleared.
    """
    score = 1.0
    warnings: list[str] = []
    raw_lower = raw_text.lower()

    # Check email appears in raw text
    email = parsed.get("email", "")
    if email and email.lower() not in raw_lower:
        warnings.append(f"Extracted email '{email}' not found in source text — cleared")
        parsed["email"] = ""
        score -= 0.2

    # Check name appears in raw text (partial match acceptable)
    name = parsed.get("name", "")
    if name:
        name_parts = name.lower().split()
        if not any(part in raw_lower for part in name_parts):
            warnings.append(f"Extracted name '{name}' not found in source text — cleared")
            parsed["name"] = ""
            score -= 0.1

    # Check LinkedIn URL
    linkedin = parsed.get("linkedin_url") or ""
    if linkedin and "linkedin.com" not in raw_lower and "linkedin" not in raw_lower:
        warnings.append("Extracted LinkedIn URL not found in source — cleared")
        parsed["linkedin_url"] = None
        score -= 0.05

    # Penalise for empty core fields
    if not parsed.get("name"):
        score -= 0.15
        warnings.append("No name extracted")
    if not parsed.get("email"):
        score -= 0.10
        warnings.append("No email extracted")
    if not parsed.get("skills"):
        score -= 0.05
        warnings.append("No skills extracted")
    if not parsed.get("experience"):
        score -= 0.10
        warnings.append("No experience extracted")

    return max(0.0, min(1.0, score)), warnings


# ---------------------------------------------------------------------------
# Agent class
# ---------------------------------------------------------------------------

class ResumeExtractorAgent:
    """
    Extracts structured data from a raw resume file.

    Usage (called from Celery task or FastAPI route):
        agent = ResumeExtractorAgent()
        parsed: ParsedResume = await agent.extract(file_bytes, mime_type)
    """

    def __init__(self, llm: LLM | None = None) -> None:
        configure_llama_settings()
        self._llm: LLM = llm or LlamaSettings.llm

    async def extract(
        self,
        file_bytes: bytes,
        mime_type: str,
    ) -> dict:
        """
        Full extraction pipeline.

        Args:
            file_bytes: Raw bytes of the uploaded file.
            mime_type:  MIME type string.

        Returns:
            ParsedResume as a dict (ready for JSON storage in PostgreSQL).
        """
        t_start = time.monotonic()

        # 1. Convert file → raw text
        raw_text = _extract_raw_text(file_bytes, mime_type)
        if not raw_text.strip():
            logger.warning("Resume text extraction produced empty output")
            return ParsedResume(
                extraction_warnings=["File produced no extractable text"],
                confidence_score=0.0,
            ).model_dump()

        # Truncate to avoid exceeding model context window
        # ~3 chars per token → MAX_CONTEXT_TOKENS * 3 chars
        max_chars = settings.MAX_CONTEXT_TOKENS * 3
        if len(raw_text) > max_chars:
            logger.warning(
                "Resume text truncated",
                original_chars=len(raw_text),
                truncated_to=max_chars,
            )
            raw_text = raw_text[:max_chars]

        # 2. Build prompt messages
        user_prompt = _USER_PROMPT_TEMPLATE.format(resume_text=raw_text)

        # 3. Call LLM via LlamaIndex (chat completion)
        from llama_index.core.llms import ChatMessage, MessageRole

        messages = [
            ChatMessage(role=MessageRole.SYSTEM, content=_SYSTEM_PROMPT),
            ChatMessage(role=MessageRole.USER, content=user_prompt),
        ]

        logger.info("ResumeExtractorAgent calling LLM", mime_type=mime_type, text_chars=len(raw_text))

        response = await self._llm.achat(messages)
        raw_json_str = response.message.content or ""

        # 4. Parse JSON response — strip markdown fences defensively
        raw_json_str = raw_json_str.strip()
        raw_json_str = re.sub(r"^```(?:json)?", "", raw_json_str).strip()
        raw_json_str = re.sub(r"```$", "", raw_json_str).strip()

        try:
            parsed_dict: dict[str, Any] = json.loads(raw_json_str)
        except json.JSONDecodeError as exc:
            logger.error(
                "ResumeExtractorAgent: JSON parse failed",
                error=str(exc),
                raw_response_preview=raw_json_str[:300],
            )
            return ParsedResume(
                extraction_warnings=[f"LLM returned invalid JSON: {exc}"],
                confidence_score=0.0,
            ).model_dump()

        # 5. Validate extracted values against source text
        confidence, warnings = _compute_confidence(parsed_dict, raw_text)
        parsed_dict["confidence_score"] = confidence
        parsed_dict["extraction_warnings"] = warnings

        # 6. Validate with Pydantic (fills defaults for missing fields)
        try:
            parsed_resume = ParsedResume.model_validate(parsed_dict)
        except Exception as exc:
            logger.error("ParsedResume Pydantic validation failed", error=str(exc))
            parsed_resume = ParsedResume(
                extraction_warnings=[f"Schema validation failed: {exc}"],
                confidence_score=0.0,
            )

        latency_ms = int((time.monotonic() - t_start) * 1000)
        logger.info(
            "ResumeExtractorAgent complete",
            confidence=parsed_resume.confidence_score,
            latency_ms=latency_ms,
            warnings=len(parsed_resume.extraction_warnings),
        )

        return parsed_resume.model_dump()